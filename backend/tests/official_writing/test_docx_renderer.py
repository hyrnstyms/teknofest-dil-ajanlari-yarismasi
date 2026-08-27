"""
backend/tests/official_writing/test_docx_renderer.py
────────────────────────────────────────────────────────────────────────────
DOCX Renderer testleri.

render_to_docx() fonksiyonunun biçimlendirilmiş .docx dosyası ürettiğini
doğrular. Sadece "hata vermedi" ile yetinilmez; üretilen belge
python-docx ile geri okunup alanların metinde gerçekten geçtiği,
sayfa düzeni parametrelerinin sayısal olarak doğru olduğu kontrol edilir.
"""

from __future__ import annotations

import io
from copy import deepcopy

import pytest
from docx import Document
from docx.shared import Mm, Cm

from backend.app.official_writing.docx_renderer import render_to_docx


# ── Yardımcı: Tüm paragraf metinlerini toplar ─────────────────────────────

def _all_text(doc: Document) -> str:
    """Belgedeki tüm paragraf metinlerini tek string olarak birleştirir."""
    return "\n".join(p.text for p in doc.paragraphs)


# ── Ortak fixture: Minimal üst yazı context'i ──────────────────────────────

@pytest.fixture()
def ust_yazi_context() -> dict:
    """build_official_writing_context() çıktısının 'context' anahtarına
    karşılık gelen minimal üst yazı context'i."""
    return {
        "tc_baslik": {
            "idare_adi": "ISPARTA KAYMAKAMLIĞI",
            "birim_adi": "Yazı İşleri Müdürlüğü",
        },
        "sayi": "E-12345678-100.01-001",
        "tarih": "24.08.2026",
        "konu": "Personel Devamsızlık Bildirimi",
        "muhatap": {
            "tur": "kurum",
            "isim": "ISPARTA VALİLİĞİ",
        },
        "muhatap_turu": "kurum_ust",
        "kapalis_ifadesi": "arz ederim.",
        "ilgi": [],
        "metin_paragraflari": [
            "Biriminizde görevli personelin devamsızlık durumu hakkında bilgi verilmesi gerekmektedir.",
            "Konuyla ilgili gerekli işlemlerin yapılması hususunda gereğini arz ederim.",
        ],
        "imza": {
            "ad_soyad": "Mehmet YILMAZ",
            "unvan": "Kaymakam",
            "yetki_turu": "normal",
        },
        "ekler": [],
        "dagitim": None,
        "iletisim": {"adres": "", "irtibat": ""},
        "sayfa_no": None,
        "uygunsuz_belge_uyarisi": None,
    }


@pytest.fixture()
def cevap_yazisi_context() -> dict:
    """İlgi zorunlu cevap yazısı context'i."""
    return {
        "tc_baslik": {
            "idare_adi": "ISPARTA KAYMAKAMLIĞI",
            "birim_adi": "Yazı İşleri Müdürlüğü",
        },
        "sayi": "E-12345678-100.01-002",
        "tarih": "24.08.2026",
        "konu": "Bilgi Talebi Cevabı",
        "muhatap": {
            "tur": "kurum",
            "isim": "BURDUR VALİLİĞİ",
        },
        "muhatap_turu": "kurum_ust",
        "kapalis_ifadesi": "arz ederim.",
        "ilgi": [
            {
                "tarih": "10.07.2026",
                "sayi": "E-98765432-200.02-555",
                "aciklama": "ilgi yazınız",
            }
        ],
        "metin_paragraflari": [
            "İlgi yazınızda belirtilen husus incelenmiştir.",
            "Gerekli bilgiler ekte sunulmuştur.",
        ],
        "imza": {
            "ad_soyad": "Ayşe KARA",
            "unvan": "Vali Yardımcısı",
            "yetki_turu": "normal",
        },
        "ekler": [
            {"ad": "Bilgi Notu", "bilgi": "3 sayfa"},
        ],
        "dagitim": None,
        "iletisim": {"adres": "Isparta Merkez", "irtibat": "Yazı İşleri"},
        "sayfa_no": None,
        "uygunsuz_belge_uyarisi": None,
    }


# ── Test 1: Üst Yazı — Temel İçerik Kontrolü ──────────────────────────────

class TestRenderToDocxUstYazi:
    """Üst yazı DOCX üretiminin içerik doğruluğu."""

    def test_produces_valid_docx(self, ust_yazi_context):
        """render_to_docx() geçerli bir .docx BytesIO döndürür."""
        result = render_to_docx(ust_yazi_context)
        assert isinstance(result, io.BytesIO)
        # Geri okuyabilmeli — bozuk dosya ise Document() hata verir
        doc = Document(result)
        assert len(doc.paragraphs) > 0

    def test_contains_tc_baslik(self, ust_yazi_context):
        """T.C. başlık bloğu belgede geçmeli."""
        result = render_to_docx(ust_yazi_context)
        doc = Document(result)
        text = _all_text(doc)
        assert "T.C." in text
        assert "ISPARTA KAYMAKAMLIĞI" in text
        assert "Yazı İşleri Müdürlüğü" in text

    def test_contains_sayi_and_tarih(self, ust_yazi_context):
        """Sayı ve Tarih belgede geçmeli."""
        result = render_to_docx(ust_yazi_context)
        doc = Document(result)
        text = _all_text(doc)
        assert "E-12345678-100.01-001" in text
        assert "24.08.2026" in text

    def test_contains_konu(self, ust_yazi_context):
        """Konu alanı belgede geçmeli."""
        result = render_to_docx(ust_yazi_context)
        doc = Document(result)
        text = _all_text(doc)
        assert "Personel Devamsızlık Bildirimi" in text

    def test_contains_muhatap(self, ust_yazi_context):
        """Muhatap belgede geçmeli."""
        result = render_to_docx(ust_yazi_context)
        doc = Document(result)
        text = _all_text(doc)
        assert "ISPARTA VALİLİĞİ" in text

    def test_contains_metin_paragraflari(self, ust_yazi_context):
        """Metin paragrafları belgede geçmeli."""
        result = render_to_docx(ust_yazi_context)
        doc = Document(result)
        text = _all_text(doc)
        assert "devamsızlık durumu" in text
        assert "gereğini arz ederim" in text

    def test_contains_kapalis_ifadesi(self, ust_yazi_context):
        """Kapanış ifadesi belgede geçmeli."""
        result = render_to_docx(ust_yazi_context)
        doc = Document(result)
        text = _all_text(doc)
        assert "arz ederim." in text

    def test_contains_imza_block(self, ust_yazi_context):
        """İmza bloğu (ad soyad + unvan) belgede geçmeli."""
        result = render_to_docx(ust_yazi_context)
        doc = Document(result)
        text = _all_text(doc)
        assert "Mehmet YILMAZ" in text
        assert "Kaymakam" in text


# ── Test 2: Cevap Yazısı — İlgi + Ek Kontrolü ─────────────────────────────

class TestRenderToDocxCevapYazisi:
    """Cevap yazısı DOCX üretiminin içerik doğruluğu."""

    def test_contains_ilgi(self, cevap_yazisi_context):
        """İlgi satırı belgede geçmeli."""
        result = render_to_docx(cevap_yazisi_context)
        doc = Document(result)
        text = _all_text(doc)
        assert "10.07.2026" in text
        assert "E-98765432-200.02-555" in text
        assert "ilgi yazınız" in text

    def test_contains_kapalis_arz_ederim(self, cevap_yazisi_context):
        """Kurum üst makam kapanışı 'arz ederim.' olmalı."""
        result = render_to_docx(cevap_yazisi_context)
        doc = Document(result)
        text = _all_text(doc)
        assert "arz ederim." in text

    def test_contains_ek(self, cevap_yazisi_context):
        """Ek listesi belgede geçmeli."""
        result = render_to_docx(cevap_yazisi_context)
        doc = Document(result)
        text = _all_text(doc)
        assert "Bilgi Notu" in text
        assert "3 sayfa" in text

    def test_contains_iletisim(self, cevap_yazisi_context):
        """İletişim bilgisi belgede geçmeli."""
        result = render_to_docx(cevap_yazisi_context)
        doc = Document(result)
        text = _all_text(doc)
        assert "Isparta Merkez" in text

    def test_contains_imza(self, cevap_yazisi_context):
        """İmza bloğu belgede geçmeli."""
        result = render_to_docx(cevap_yazisi_context)
        doc = Document(result)
        text = _all_text(doc)
        assert "Ayşe KARA" in text
        assert "Vali Yardımcısı" in text


# ── Test 3: Sayfa Düzeni — Sayısal Doğrulama ──────────────────────────────

class TestRenderToDocxPageLayout:
    """Sayfa düzeni parametrelerinin sayısal olarak doğruluğu.

    python-docx dahili olarak EMU (English Metric Units) kullanır.
    Mm/Cm dönüşümlerinde küçük yuvarlama farkları olabilir;
    bu nedenle tolerans payı bırakılır (±1000 EMU ≈ ±0.1 mm).
    """

    EMU_TOLERANCE = 5000  # ~0.5 mm tolerans

    def test_page_width_is_a4(self, ust_yazi_context):
        """Sayfa genişliği A4 = 210 mm olmalı (±0.5mm tolerans)."""
        result = render_to_docx(ust_yazi_context)
        doc = Document(result)
        section = doc.sections[0]
        expected = Mm(210)
        actual = section.page_width
        assert abs(actual - expected) <= self.EMU_TOLERANCE, (
            f"Sayfa genişliği {actual} EMU, beklenen {expected} EMU "
            f"(fark: {abs(actual - expected)} EMU, tolerans: {self.EMU_TOLERANCE} EMU)"
        )

    def test_page_height_is_a4(self, ust_yazi_context):
        """Sayfa yüksekliği A4 = 297 mm olmalı (±0.5mm tolerans)."""
        result = render_to_docx(ust_yazi_context)
        doc = Document(result)
        section = doc.sections[0]
        expected = Mm(297)
        actual = section.page_height
        assert abs(actual - expected) <= self.EMU_TOLERANCE, (
            f"Sayfa yüksekliği {actual} EMU, beklenen {expected} EMU "
            f"(fark: {abs(actual - expected)} EMU, tolerans: {self.EMU_TOLERANCE} EMU)"
        )

    def test_top_margin_is_1_5_cm(self, ust_yazi_context):
        """Üst kenar boşluğu 1,5 cm olmalı (±0.5mm tolerans)."""
        result = render_to_docx(ust_yazi_context)
        doc = Document(result)
        section = doc.sections[0]
        expected = Cm(1.5)
        actual = section.top_margin
        assert abs(actual - expected) <= self.EMU_TOLERANCE, (
            f"Üst kenar boşluğu {actual} EMU, beklenen {expected} EMU "
            f"(fark: {abs(actual - expected)} EMU, tolerans: {self.EMU_TOLERANCE} EMU)"
        )

    def test_bottom_margin_is_1_5_cm(self, ust_yazi_context):
        """Alt kenar boşluğu 1,5 cm olmalı (±0.5mm tolerans)."""
        result = render_to_docx(ust_yazi_context)
        doc = Document(result)
        section = doc.sections[0]
        expected = Cm(1.5)
        actual = section.bottom_margin
        assert abs(actual - expected) <= self.EMU_TOLERANCE, (
            f"Alt kenar boşluğu {actual} EMU, beklenen {expected} EMU "
            f"(fark: {abs(actual - expected)} EMU, tolerans: {self.EMU_TOLERANCE} EMU)"
        )

    def test_left_margin_is_1_5_cm(self, ust_yazi_context):
        """Sol kenar boşluğu 1,5 cm olmalı (±0.5mm tolerans)."""
        result = render_to_docx(ust_yazi_context)
        doc = Document(result)
        section = doc.sections[0]
        expected = Cm(1.5)
        actual = section.left_margin
        assert abs(actual - expected) <= self.EMU_TOLERANCE, (
            f"Sol kenar boşluğu {actual} EMU, beklenen {expected} EMU "
            f"(fark: {abs(actual - expected)} EMU, tolerans: {self.EMU_TOLERANCE} EMU)"
        )

    def test_right_margin_is_1_5_cm(self, ust_yazi_context):
        """Sağ kenar boşluğu 1,5 cm olmalı (±0.5mm tolerans)."""
        result = render_to_docx(ust_yazi_context)
        doc = Document(result)
        section = doc.sections[0]
        expected = Cm(1.5)
        actual = section.right_margin
        assert abs(actual - expected) <= self.EMU_TOLERANCE, (
            f"Sağ kenar boşluğu {actual} EMU, beklenen {expected} EMU "
            f"(fark: {abs(actual - expected)} EMU, tolerans: {self.EMU_TOLERANCE} EMU)"
        )

    def test_font_is_times_new_roman_12pt(self, ust_yazi_context):
        """İlk içerikli paragrafın yazı tipi Times New Roman 12pt olmalı."""
        result = render_to_docx(ust_yazi_context)
        doc = Document(result)
        # İlk paragraf "T.C." olmalı
        first_para = doc.paragraphs[0]
        assert first_para.text.strip() == "T.C."
        run = first_para.runs[0]
        assert run.font.name == "Times New Roman"
        from docx.shared import Pt
        assert run.font.size == Pt(12)


# ── Test 4: QR Doğrulama Kodu ─────────────────────────────────────────────

class TestRenderToDocxQRCode:
    """QR doğrulama kodu ekleme/yoklama testleri."""

    SAMPLE_EVRAK_ID = "a1b2c3d4-e5f6-7890-abcd-ef1234567890"

    def test_render_to_docx_with_qr_contains_image(self, ust_yazi_context):
        """evrak_id verildiğinde DOCX en az bir inline shape (görüntü) içermeli."""
        result = render_to_docx(ust_yazi_context, evrak_id=self.SAMPLE_EVRAK_ID)
        doc = Document(result)
        # python-docx: inline_shapes belgedeki tüm inline görüntüleri listeler
        assert len(doc.inline_shapes) >= 1, (
            "evrak_id verildiğinde DOCX'te en az bir inline shape (QR görüntüsü) bekleniyor."
        )

    def test_render_to_docx_without_evrak_id_no_qr(self, ust_yazi_context):
        """evrak_id verilmezse DOCX'te inline shape (QR) olmamalı."""
        result = render_to_docx(ust_yazi_context)
        doc = Document(result)
        assert len(doc.inline_shapes) == 0, (
            "evrak_id verilmezse DOCX'te inline shape olmamalı (geriye dönük uyumluluk)."
        )

    def test_render_to_docx_qr_content_is_decodable(self, ust_yazi_context):
        """QR kodunun içeriği cv2 ile decode edilip evrak_id ile eşleşmeli."""
        import cv2
        import numpy as np
        from docx.opc.constants import RELATIONSHIP_TYPE as RT

        result = render_to_docx(ust_yazi_context, evrak_id=self.SAMPLE_EVRAK_ID)
        doc = Document(result)

        # DOCX'ten görüntü blob'larını çıkar
        image_blobs = []
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_blobs.append(rel.target_part.blob)

        assert len(image_blobs) >= 1, "DOCX'te en az bir görüntü bekleniyor."

        # İlk görüntüyü cv2 ile decode et
        img_bytes = image_blobs[0]
        arr = np.frombuffer(img_bytes, np.uint8)
        img = cv2.imdecode(arr, cv2.IMREAD_GRAYSCALE)
        assert img is not None, "Görüntü cv2 ile okunamadı."

        detector = cv2.QRCodeDetector()
        decoded_value, _, _ = detector.detectAndDecode(img)

        import os
        base_url = os.environ.get("KAMUAI_BASE_URL", "http://localhost:8000").rstrip("/")
        expected = f"{base_url}/api/verify/{self.SAMPLE_EVRAK_ID}"
        assert decoded_value == expected, (
            f"QR içeriği eşleşmiyor. Beklenen: '{expected}', "
            f"Okunan: '{decoded_value}'"
        )

    def test_qr_dogrulama_text_present(self, ust_yazi_context):
        """QR eklendiğinde doğrulama açıklama metni de belgede olmalı."""
        result = render_to_docx(ust_yazi_context, evrak_id=self.SAMPLE_EVRAK_ID)
        doc = Document(result)
        text = _all_text(doc)
        assert "QR kodu okutunuz" in text


# ── Test 5: Endpoint Entegrasyon — analysis_id QR'a ulaşıyor mu ────────────

class TestExportDocxEndpointQR:
    """FastAPI endpoint'inin analysis_id'yi QR'a doğru geçirdiğini doğrular."""

    def test_export_endpoint_passes_analysis_id_to_qr(self):
        """GET /api/analysis/{id}/export/docx endpoint'i analysis_id'yi
        QR koduna doğru şekilde kodlamalı."""
        import cv2
        import numpy as np
        from fastapi.testclient import TestClient
        from backend.app.main import app, analysis_store

        client = TestClient(app)

        # Sahte analiz kaydı oluştur
        test_id = "test-uuid-1234-5678-abcdef012345"
        analysis_store[test_id] = {
            "draft": {
                "body": "Test metni.",
                "subject": "Test Konusu",
            },
            "draft_type": "ust_yazi",
            "extraction": {},
            "routing": {},
            "kurum_profili_id": "kaymakamlik_v1",
        }

        try:
            response = client.get(f"/api/analysis/{test_id}/export/docx")
            assert response.status_code == 200
            assert "wordprocessingml" in response.headers.get("content-type", "")

            # DOCX'i geri oku
            docx_bytes = io.BytesIO(response.content)
            doc = Document(docx_bytes)

            # En az bir inline shape (QR) olmalı
            assert len(doc.inline_shapes) >= 1, (
                "Endpoint üzerinden üretilen DOCX'te QR görüntüsü bekleniyor."
            )

            # QR'ı decode et ve analysis_id'nin doğru kodlandığını kontrol et
            image_blobs = []
            for rel in doc.part.rels.values():
                if "image" in rel.reltype:
                    image_blobs.append(rel.target_part.blob)

            assert len(image_blobs) >= 1

            arr = np.frombuffer(image_blobs[0], np.uint8)
            img = cv2.imdecode(arr, cv2.IMREAD_GRAYSCALE)
            assert img is not None

            detector = cv2.QRCodeDetector()
            decoded_value, _, _ = detector.detectAndDecode(img)

            import os
            base_url = os.environ.get("KAMUAI_BASE_URL", "http://localhost:8000").rstrip("/")
            expected = f"{base_url}/api/verify/{test_id}"
            assert decoded_value == expected, (
                f"Endpoint QR içeriği eşleşmiyor. "
                f"Beklenen: '{expected}', Okunan: '{decoded_value}'"
            )
        finally:
            # Temizle
            analysis_store.pop(test_id, None)

    def test_export_uses_mod_c_validated_context_without_rebuilding(
        self,
        ust_yazi_context,
        monkeypatch,
    ):
        """Mod C konusu, extraction konusu tarafından DOCX'te geri alınmamalı."""
        from fastapi.testclient import TestClient
        from backend.app.main import app, analysis_store

        client = TestClient(app)
        test_id = "mod-c-context-export-test"
        candidate_context = deepcopy(ust_yazi_context)
        candidate_context["konu"] = "Düzenlenmiş Başvuru Konusu"

        analysis_store[test_id] = {
            "draft": {
                "draft_type": "ust_yazi",
                "draft": {
                    "subject": "Düzenlenmiş Başvuru Konusu",
                    "body": "Başvuru işlemi tamamlanmıştır.",
                },
                "mod_c_validated_context": candidate_context,
            },
            "extraction": {
                "fields": {
                    "subject": {
                        "value": "Eski Başvuru Konusu",
                        "validated": True,
                    }
                }
            },
            "routing": {},
            "kurum_profili_id": "kaymakamlik_v1",
        }

        def fail_if_adapter_is_called(*args, **kwargs):
            raise AssertionError("Mod C context varken adapter çağrılmamalı")

        monkeypatch.setattr(
            "backend.app.official_writing.context_adapter."
            "build_official_writing_context",
            fail_if_adapter_is_called,
        )

        try:
            response = client.get(f"/api/analysis/{test_id}/export/docx")
            assert response.status_code == 200

            doc = Document(io.BytesIO(response.content))
            text = _all_text(doc)
            assert "Düzenlenmiş Başvuru Konusu" in text
            assert "Eski Başvuru Konusu" not in text
        finally:
            analysis_store.pop(test_id, None)

