from pathlib import Path

from docx import Document as DocxDocument


def load_docx(path: Path) -> str:
    """
    DOCX dosyasındaki paragrafları ve tabloları metne dönüştürür.
    """

    document = DocxDocument(path)

    parts = []

    # Paragraflar
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()

        if text:
            parts.append(text)

    # Tablolar
    for table in document.tables:
        for row in table.rows:
            cells = []

            for cell in row.cells:
                text = cell.text.strip()

                if text:
                    cells.append(text)

            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts).strip()