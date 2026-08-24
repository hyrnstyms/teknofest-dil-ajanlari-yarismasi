"""
backend/app/official_writing/docx_renderer.py
──────────────────────────────────────────────────────────────────────────────
Resmî Yazı DOCX Üretici

build_official_writing_context() tarafından üretilen yapılandırılmış
context sözlüğünü kullanarak, biçim kurallarına uygun bir .docx dosyası
üretir.

Sayfa düzeni kuralları (kaynak: docs/format_kurallari_checklist.md):
  - Kağıt: A4 (210×297 mm)
  - Kenar boşluğu: üst/sol/sağ/alt 1,5 cm
  - Yazı tipi: Times New Roman 12pt
  - Satır aralığı: 1.0 (tek satır) [TASARIM KARARI]
  - Paragraf girintisi: 1,25 cm
  - İmza bloğu: sağda ortalı
  - Sayı/Tarih: aynı satırda, sayı solda, tarih sağda (tab stop ile)

Bu modül mevcut format_validator.py ve template_renderer.py'ye DOKUNMAZ;
sadece EK bir çıktı formatı sağlar.
"""

from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def _set_run_font(run, name: str = "Times New Roman", size_pt: int = 12, bold: bool = False):
    """Bir run'a yazı tipi, boyut ve kalınlık ayarı uygular."""
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    # Times New Roman için rFonts ayarı
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = run._element.makeelement(qn("w:rFonts"), {})
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:cs"), name)


def _set_paragraph_spacing(paragraph, line_spacing: float = 1.0):
    """Paragrafın satır aralığını ve paragraf öncesi/sonrası boşluğunu ayarlar."""
    pf = paragraph.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)


def _add_paragraph(
    doc: Document,
    text: str,
    alignment=WD_ALIGN_PARAGRAPH.LEFT,
    bold: bool = False,
    font_size: int = 12,
    first_line_indent: float | None = None,
    left_indent: float | None = None,
) -> Any:
    """Belgeye standart biçimli bir paragraf ekler."""
    p = doc.add_paragraph()
    p.alignment = alignment
    run = p.add_run(text)
    _set_run_font(run, size_pt=font_size, bold=bold)
    _set_paragraph_spacing(p)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    if left_indent is not None:
        p.paragraph_format.left_indent = Cm(left_indent)
    return p


def _add_empty_paragraph(doc: Document) -> Any:
    """Boş bir satır ekler."""
    p = doc.add_paragraph()
    run = p.add_run("")
    _set_run_font(run)
    _set_paragraph_spacing(p)
    return p


def _add_sayi_tarih_paragraph(doc: Document, sayi: str, tarih: str) -> Any:
    """
    Sayı ve Tarih'i aynı satırda gösterir.
    Sayı sol hizalı, Tarih sağa hizalı (tab stop ile).
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_paragraph_spacing(p)

    # Sağa hizalı tab stop ekle (sayfa genişliği - kenar boşlukları = ~17 cm)
    pPr = p._element.get_or_add_pPr()
    tabs = pPr.find(qn("w:tabs"))
    if tabs is None:
        tabs = p._element.makeelement(qn("w:tabs"), {})
        pPr.append(tabs)
    tab_elem = p._element.makeelement(qn("w:tab"), {
        qn("w:val"): "right",
        qn("w:pos"): str(int(Cm(17.0))),
    })
    tabs.append(tab_elem)

    # Sayı run'u
    run_sayi = p.add_run("Say\u0131: " + sayi)
    _set_run_font(run_sayi)

    # Tab karakteri
    run_tab = p.add_run("\t")
    _set_run_font(run_tab)

    # Tarih run'u
    run_tarih = p.add_run("Tarih: " + tarih)
    _set_run_font(run_tarih)

    return p


def _add_imza_block(doc: Document, imza: dict[str, Any]) -> None:
    """İmza bloğunu sağda ortalı olarak ekler."""
    # 2 boş satır bırak
    _add_empty_paragraph(doc)
    _add_empty_paragraph(doc)

    ad_soyad = imza.get("ad_soyad", "")
    unvan = imza.get("unvan", "")
    yetki_turu = imza.get("yetki_turu", "normal")
    vekil_makam = imza.get("vekil_makam", "")

    # İmza satırları — sağda konumlandırılmış, ortalı
    indent_cm = 10.0

    _add_paragraph(doc, ad_soyad, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                   left_indent=indent_cm)

    if yetki_turu == "yetki_devri" and vekil_makam:
        _add_paragraph(doc, vekil_makam + " a.",
                       alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       left_indent=indent_cm)
    elif yetki_turu == "vekaletname" and vekil_makam:
        _add_paragraph(doc, vekil_makam + " V.",
                       alignment=WD_ALIGN_PARAGRAPH.CENTER,
                       left_indent=indent_cm)

    _add_paragraph(doc, unvan, alignment=WD_ALIGN_PARAGRAPH.CENTER,
                   left_indent=indent_cm)


def render_to_docx(context: dict[str, Any]) -> io.BytesIO:
    """
    build_official_writing_context() çıktısındaki 'context' sözlüğünü
    kullanarak biçimlendirilmiş bir .docx dosyası üretir.

    Args:
        context: build_official_writing_context()["context"] sözlüğü.

    Returns:
        BytesIO nesnesi (.docx içeriği).
    """
    doc = Document()

    # ── Sayfa Düzeni ────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)

    # ── 1. BAŞLIK BLOĞU (Madde 10) ─────────────────────────────────────────
    _add_paragraph(doc, "T.C.", alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)

    tc_baslik = context.get("tc_baslik", {})
    idare_adi = tc_baslik.get("idare_adi", "").upper()
    birim_adi = tc_baslik.get("birim_adi", "")

    _add_paragraph(doc, idare_adi, alignment=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    _add_paragraph(doc, birim_adi, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Başlıktan sonra boş satır
    _add_empty_paragraph(doc)

    # ── 2. SAYI ve TARİH (Madde 11, 12) ───────────────────────────────────
    sayi = context.get("sayi", "")
    tarih = context.get("tarih", "")
    _add_sayi_tarih_paragraph(doc, sayi, tarih)

    # ── 3. KONU (Madde 13) ─────────────────────────────────────────────────
    konu = context.get("konu", "")
    _add_paragraph(doc, "Konu: " + konu, bold=True)

    # Konudan sonra 2 boş satır
    _add_empty_paragraph(doc)

    # ── 4. MUHATAP (Madde 14) ──────────────────────────────────────────────
    muhatap = context.get("muhatap", {})
    muhatap_tur = muhatap.get("tur", "")
    muhatap_isim = muhatap.get("isim", "")

    if muhatap_tur == "kurum":
        muhatap_text = muhatap_isim.upper()
    elif muhatap_tur == "gercek_kisi":
        muhatap_text = "Say\u0131n " + muhatap_isim
    elif muhatap_tur == "dagitim":
        muhatap_text = "DA\u011eITIM YERLER\u0130NE"
    else:
        muhatap_text = muhatap_isim

    _add_paragraph(doc, muhatap_text, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Muhataptan sonra 2 boş satır
    _add_empty_paragraph(doc)

    # ── 5. İLGİ (Madde 15) ─────────────────────────────────────────────────
    ilgi = context.get("ilgi", [])
    if ilgi:
        for ilgi_item in ilgi:
            ilgi_tarih = ilgi_item.get("tarih", "")
            ilgi_sayi = ilgi_item.get("sayi", "")
            ilgi_aciklama = ilgi_item.get("aciklama", "")
            ilgi_text = (
                "\u0130lgi: " + ilgi_tarih + " tarihli ve "
                + ilgi_sayi + " say\u0131l\u0131 " + ilgi_aciklama + "."
            )
            _add_paragraph(doc, ilgi_text)
        _add_empty_paragraph(doc)

    # ── 6. METİN / GÖVDE (Madde 16) ───────────────────────────────────────
    metin_paragraflari = context.get("metin_paragraflari", [])
    for paragraf in metin_paragraflari:
        _add_paragraph(
            doc, paragraf,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_line_indent=1.25,
        )

    # ── 7. KAPANIŞ İFADESİ (Madde 16/12) ──────────────────────────────────
    kapalis = context.get("kapalis_ifadesi", "")
    if kapalis:
        _add_empty_paragraph(doc)
        _add_paragraph(
            doc, kapalis,
            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
            first_line_indent=1.25,
        )

    # ── 8. İMZA BLOĞU (Madde 17) ──────────────────────────────────────────
    imza = context.get("imza", {})
    if imza:
        _add_imza_block(doc, imza)

    # ── 9. EK LİSTESİ (Madde 18) ──────────────────────────────────────────
    ekler = context.get("ekler", [])
    if ekler:
        _add_empty_paragraph(doc)
        _add_empty_paragraph(doc)
        if len(ekler) == 1:
            ek = ekler[0]
            _add_paragraph(doc, "Ek: " + ek.get("ad", "") + " (" + ek.get("bilgi", "") + ")")
        else:
            _add_paragraph(doc, "Ek: " + str(len(ekler)) + " adet")
            for i, ek in enumerate(ekler, 1):
                _add_paragraph(doc, str(i) + ". " + ek.get("ad", "") + " (" + ek.get("bilgi", "") + ")")

    # ── 10. DAĞITIM LİSTESİ (Madde 19) ────────────────────────────────────
    dagitim = context.get("dagitim")
    if dagitim:
        _add_empty_paragraph(doc)
        _add_empty_paragraph(doc)
        _add_paragraph(doc, "Da\u011f\u0131t\u0131m:")
        geregi = dagitim.get("geregi", [])
        bilgi = dagitim.get("bilgi", [])
        if geregi:
            _add_paragraph(doc, "Gere\u011fi:", bold=True)
            for yer in geregi:
                _add_paragraph(doc, yer)
        if bilgi:
            _add_paragraph(doc, "Bilgi:", bold=True)
            for yer in bilgi:
                _add_paragraph(doc, yer)

    # ── 11. İLETİŞİM BİLGİSİ (Madde 10) ──────────────────────────────────
    iletisim = context.get("iletisim", {})
    adres = iletisim.get("adres", "")
    irtibat = iletisim.get("irtibat", "")
    if adres or irtibat:
        _add_empty_paragraph(doc)
        _add_paragraph(doc, "_" * 80)
        iletisim_parts = []
        if adres:
            iletisim_parts.append(adres)
        if irtibat:
            iletisim_parts.append(irtibat)
        telefon = iletisim.get("telefon", "")
        if telefon:
            iletisim_parts.append(telefon)
        _add_paragraph(doc, " | ".join(iletisim_parts), font_size=10)

    # ── Kaydet ─────────────────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
